# Script para leer las plantillas del diplomado
from docx import Document
import os

carpeta = "docs/luego elimino no subir a github"

# Leer plantilla de resúmenes
print("="*60)
print("PLANTILLA DE RESUMENES")
print("="*60)
try:
    doc = Document(f"{carpeta}/Plantilla_Resumenes_EventoDiplomado.docx")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
except Exception as e:
    print(f"Error: {e}")

print("\n")

# Leer informe anterior
print("="*60)
print("INFORME ANTERIOR (para referencia)")
print("="*60)
try:
    doc = Document(f"{carpeta}/paez_jean_Actividad 3 Desarrollo de prototipo con contexto Personalizado.docx")
    for i, para in enumerate(doc.paragraphs[:50]):  # Primeros 50 párrafos
        if para.text.strip():
            print(para.text[:200])
except Exception as e:
    print(f"Error: {e}")

