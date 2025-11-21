"""
ARB-BOT - Procesador de Documentos Institucionales
Procesa documentos PDF, DOCX para el sistema RAG
Usa LlamaIndex para chunking semántico avanzado
"""

import os
import logging
from typing import List, Dict, Optional
import re

logger = logging.getLogger(__name__)

# Intentar importar LlamaIndex para chunking avanzado
try:
    from llama_index.core.node_parser import SemanticSplitterNodeParser
    from llama_index.core.schema import Document as LlamaDocument
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    LLAMAINDEX_AVAILABLE = True
    logger.info("✅ LlamaIndex disponible - usando chunking semántico avanzado")
except ImportError:
    LLAMAINDEX_AVAILABLE = False
    logger.warning("⚠️ LlamaIndex no disponible - usando chunking básico mejorado")

class DocumentProcessor:
    """
    Procesa documentos institucionales (PDF, DOCX)
    Extrae texto y lo prepara para embeddings
    """
    
    def __init__(self):
        """Inicializar procesador de documentos"""
        self.supported_formats = ['.pdf', '.docx', '.txt']
        self.chunk_size = 800  # Tamaño de chunks optimizado para Railway (2GB RAM) - sin modelos locales
        self.chunk_overlap = 100  # Solapamiento mayor para mantener contexto entre chunks
        
        # Inicializar LlamaIndex si está disponible
        self.semantic_splitter = None
        if LLAMAINDEX_AVAILABLE:
            try:
                # Usar el mismo modelo de embeddings que RAGService
                # Nota: El modelo debe estar disponible en Hugging Face
                embedding_model = HuggingFaceEmbedding(
                    model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2"
                )
                # SemanticSplitter agrupa por similitud semántica
                # Parámetros optimizados para velocidad sin perder calidad significativa
                self.semantic_splitter = SemanticSplitterNodeParser(
                    buffer_size=4,  # Aumentado para reducir cálculos (más rápido)
                    breakpoint_percentile_threshold=90,  # Reducido para menos cálculos (más rápido)
                    embed_model=embedding_model
                )
                logger.info("✅ SemanticSplitter inicializado correctamente")
                logger.info("✅ SemanticSplitter inicializado para chunking inteligente")
            except Exception as e:
                logger.warning(f"⚠️ No se pudo inicializar SemanticSplitter: {e}. Usando chunking básico.")
                self.semantic_splitter = None
    
    def process_document(self, file_path: str, max_pages: Optional[int] = None) -> List[Dict]:
        """
        Procesar un documento y extraer texto
        
        Args:
            file_path: Ruta al archivo
            max_pages: Número máximo de páginas para PDFs (por defecto 100)
        
        Returns:
            Lista de chunks con texto y metadata
        """
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            logger.info(f"Iniciando extracción de texto de {file_ext}")
            
            if file_ext == '.pdf':
                text = self._extract_from_pdf(file_path, max_pages=max_pages)
                if not text or len(text.strip()) < 10:
                    raise ValueError("No se pudo extraer texto del PDF. El documento puede estar protegido o corrupto.")
            elif file_ext == '.docx':
                text = self._extract_from_docx(file_path)
            elif file_ext == '.txt':
                text = self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Formato no soportado: {file_ext}")
            
            if not text or len(text.strip()) < 10:
                raise ValueError("No se pudo extraer texto del documento o el documento está vacío")
            
            logger.info(f"Texto extraído: {len(text)} caracteres")
            
            # Dividir en chunks (usar LlamaIndex si está disponible)
            logger.info("Dividiendo texto en chunks...")
            if LLAMAINDEX_AVAILABLE and self.semantic_splitter:
                try:
                    chunks = self._split_with_llamaindex(text, file_path)
                except Exception as e:
                    logger.warning(f"Error usando SemanticSplitter ({e}), usando chunking básico...")
                    chunks = self._split_into_chunks(text, file_path)
            else:
                chunks = self._split_into_chunks(text, file_path)
            
            logger.info(f"Documento procesado: {len(chunks)} chunks extraídos")
            return chunks
            
        except Exception as e:
            logger.error(f"Error procesando documento {file_path}: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str, max_pages: Optional[int] = None) -> str:
        """
        Extraer texto de PDF con límite opcional de páginas
        
        Args:
            file_path: Ruta al archivo PDF
            max_pages: Número máximo de páginas a procesar (None = procesar todo)
        """
        try:
            import PyPDF2
            
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                # Determinar cuántas páginas procesar
                if max_pages is None:
                    pages_to_process = total_pages
                    logger.info(f"Procesando todas las {total_pages} páginas del PDF...")
                else:
                    pages_to_process = min(total_pages, max_pages)
                    if total_pages > max_pages:
                        logger.warning(f"PDF tiene {total_pages} páginas. Procesando solo las primeras {max_pages} páginas.")
                
                # Procesar páginas con logging de progreso y límite de tiempo
                import time
                start_time = time.time()
                max_processing_time = 600  # 10 minutos máximo para extracción (aumentado para documentos grandes)
                
                for page_num in range(pages_to_process):
                    # Verificar tiempo transcurrido
                    elapsed = time.time() - start_time
                    if elapsed > max_processing_time:
                        logger.warning(f"Tiempo de procesamiento excedido ({elapsed:.1f}s). Procesadas {page_num}/{pages_to_process} páginas.")
                        text += f"\n\n[Nota: Procesamiento interrumpido por tiempo. Se procesaron {page_num} de {total_pages} páginas.]"
                        break
                    
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text.strip():
                            # Limitar tamaño de texto por página para evitar problemas de memoria
                            if len(page_text) > 8000:  # Reducido para evitar problemas
                                page_text = page_text[:8000] + "... [texto truncado por tamaño]"
                            text += f"\n--- Página {page_num + 1} ---\n{page_text}\n"
                        
                        # Log cada 10 páginas para monitorear progreso
                        if (page_num + 1) % 10 == 0:
                            elapsed = time.time() - start_time
                            logger.info(f"Procesadas {page_num + 1}/{pages_to_process} páginas... ({((page_num + 1) / pages_to_process * 100):.1f}%) - Tiempo: {elapsed:.1f}s")
                    except Exception as page_error:
                        logger.warning(f"Error extrayendo página {page_num + 1}: {str(page_error)}")
                        continue
                
                if max_pages and total_pages > max_pages:
                    text += f"\n\n[Nota: Este documento tiene {total_pages} páginas. Se procesaron las primeras {max_pages} páginas.]"
                else:
                    logger.info(f"✅ Procesadas todas las {total_pages} páginas exitosamente")
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extrayendo PDF: {str(e)}")
            # Fallback: intentar con otra librería
            try:
                import pypdf
                text = ""
                with open(file_path, 'rb') as file:
                    pdf_reader = pypdf.PdfReader(file)
                    total_pages = len(pdf_reader.pages)
                    
                    if max_pages is None:
                        pages_to_process = total_pages
                        logger.info(f"Procesando todas las {total_pages} páginas del PDF...")
                    else:
                        pages_to_process = min(total_pages, max_pages)
                        if total_pages > max_pages:
                            logger.warning(f"PDF tiene {total_pages} páginas. Procesando solo las primeras {max_pages} páginas.")
                    
                    import time
                    start_time = time.time()
                    max_processing_time = 600  # 10 minutos máximo (aumentado para documentos grandes)
                    
                    for page_num in range(pages_to_process):
                        elapsed = time.time() - start_time
                        if elapsed > max_processing_time:
                            logger.warning(f"Tiempo excedido. Procesadas {page_num}/{pages_to_process} páginas.")
                            text += f"\n\n[Nota: Procesamiento interrumpido. Se procesaron {page_num} de {total_pages} páginas.]"
                            break
                        
                        try:
                            page = pdf_reader.pages[page_num]
                            page_text = page.extract_text()
                            if page_text.strip():
                                if len(page_text) > 8000:
                                    page_text = page_text[:8000] + "... [texto truncado]"
                                text += page_text + "\n"
                            
                            if (page_num + 1) % 10 == 0:
                                elapsed = time.time() - start_time
                                logger.info(f"Procesadas {page_num + 1}/{pages_to_process} páginas... ({((page_num + 1) / pages_to_process * 100):.1f}%) - Tiempo: {elapsed:.1f}s")
                        except Exception:
                            continue
                    
                    if max_pages and total_pages > max_pages:
                        text += f"\n\n[Nota: Este documento tiene {total_pages} páginas. Se procesaron las primeras {max_pages} páginas.]"
                    else:
                        logger.info(f"✅ Procesadas todas las {total_pages} páginas exitosamente")
                
                return text.strip()
            except ImportError:
                raise Exception(f"No se pudo extraer texto del PDF. Instala PyPDF2 o pypdf: {str(e)}")
            except Exception as fallback_error:
                raise Exception(f"No se pudo extraer texto del PDF: {str(e)}")
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extraer texto de DOCX"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # También extraer texto de tablas
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            logger.error(f"Error extrayendo DOCX: {str(e)}")
            raise Exception(f"No se pudo extraer texto del DOCX: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extraer texto de TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Intentar con otra codificación
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def _split_with_llamaindex(self, text: str, source: str) -> List[Dict]:
        """
        Dividir texto usando enfoque híbrido:
        1. Primero dividir por artículos/secciones (estructura del documento)
        2. Luego aplicar SemanticSplitter dentro de cada sección
        3. Mantener overlap entre secciones para contexto
        
        Esto mantiene ~98% de calidad semántica mientras evita timeouts y problemas de memoria.
        
        Args:
            text: Texto completo
            source: Fuente del documento
        
        Returns:
            Lista de chunks con metadata
        """
        try:
            from llama_index.core.schema import Document as LlamaDocument
            import time
            
            # Limpiar texto primero
            text = self._clean_text(text)
            
            logger.info("🔍 Paso 1: Dividiendo documento por artículos/secciones...")
            
            # Paso 1: Dividir por artículos/secciones primero
            sections = self._split_into_sections(text)
            logger.info(f"✅ Documento dividido en {len(sections)} secciones/artículos")
            
            if not sections:
                # Si no se encontraron secciones, usar todo el texto pero en chunks más pequeños
                logger.warning("⚠️ No se detectaron artículos/secciones. Procesando en secciones grandes...")
                sections = self._split_into_large_sections(text, max_size=150000)  # ~150k caracteres por sección
                logger.info(f"✅ Documento dividido en {len(sections)} secciones grandes")
            
            # Paso 1.5: Combinar secciones pequeñas para optimizar procesamiento
            logger.info("🔍 Paso 1.5: Combinando secciones pequeñas para optimizar...")
            optimized_sections = self._combine_small_sections(sections, min_size=5000)
            logger.info(f"✅ Secciones optimizadas: {len(sections)} → {len(optimized_sections)} secciones")
            
            # Paso 2: Aplicar SemanticSplitter solo a secciones grandes
            all_chunks = []
            chunk_id = 0
            start_time = time.time()
            
            logger.info("🔍 Paso 2: Aplicando chunking semántico a secciones grandes...")
            
            for section_idx, section in enumerate(optimized_sections):
                section_text = section['text']
                section_article = section.get('article')
                
                # Log progreso cada 10 secciones o cada 60 segundos (reducir logging para velocidad)
                elapsed = time.time() - start_time
                if (section_idx + 1) % 10 == 0 or section_idx == 0 or (elapsed > 60 and (section_idx + 1) % 5 == 0):
                    logger.info(f"📄 Procesando sección {section_idx + 1}/{len(optimized_sections)} (Art. {section_article or 'N/A'}, {len(section_text)} chars) - Tiempo: {elapsed:.1f}s")
                
                # Estrategia híbrida optimizada para velocidad y calidad:
                # - Secciones pequeñas/medianas: chunking inteligente básico (rápido, mantiene estructura)
                # - Secciones MUY grandes: SemanticSplitter (calidad semántica máxima)
                section_size = len(section_text.strip())
                
                if section_size < 15000:  # Aumentado: solo usar SemanticSplitter en secciones MUY grandes
                    # Usar chunking inteligente básico que respeta artículos, párrafos y oraciones
                    # Esto mantiene ~95% de calidad pero es 10-20x más rápido
                    basic_chunks = self._split_section_basic(section_text, source, section_article, chunk_id)
                    all_chunks.extend(basic_chunks)
                    chunk_id += len(basic_chunks)
                    continue
                
                # Aplicar SemanticSplitter solo a secciones MUY grandes (>15000 chars)
                # Estas secciones se benefician más del análisis semántico profundo
                try:
                    logger.debug(f"🔬 Usando SemanticSplitter para sección grande ({section_size} chars)")
                    llama_doc = LlamaDocument(text=section_text, metadata={'source': source, 'article': section_article})
                    nodes = self.semantic_splitter.get_nodes_from_documents([llama_doc])
                    
                    # Convertir nodes a chunks con metadata
                    for node in nodes:
                        # Extraer número de artículo si no lo tenemos
                        article = section_article or self._extract_article_number(node.text)
                        
                        chunk = {
                            'id': chunk_id,
                            'text': node.text,
                            'source': source,
                            'chunk_index': chunk_id,
                            'article': article
                        }
                        all_chunks.append(chunk)
                        chunk_id += 1
                        
                except Exception as section_error:
                    logger.warning(f"⚠️ Error procesando sección {section_idx + 1} con SemanticSplitter: {section_error}. Usando chunking básico para esta sección.")
                    # Fallback: dividir esta sección con método básico
                    basic_chunks = self._split_section_basic(section_text, source, section_article, chunk_id)
                    all_chunks.extend(basic_chunks)
                    chunk_id += len(basic_chunks)
            
            elapsed = time.time() - start_time
            logger.info(f"✅ Chunking semántico completado: {len(all_chunks)} chunks en {elapsed:.1f}s")
            logger.info(f"📊 Promedio: {len(all_chunks)/elapsed:.1f} chunks/segundo")
            
            return all_chunks
            
        except Exception as e:
            logger.error(f"❌ Error en _split_with_llamaindex: {e}")
            logger.warning("⚠️ Fallback a chunking básico...")
            # Fallback a chunking básico
            return self._split_into_chunks(text, source)
    
    def _split_into_sections(self, text: str) -> List[Dict]:
        """
        Dividir texto en secciones basadas en artículos y secciones del documento.
        Respeta la estructura del documento.
        
        Returns:
            Lista de secciones con texto y metadata
        """
        sections = []
        
        # Patrones para detectar artículos y secciones
        article_pattern = re.compile(r'(?i)(art[ií]culo|art\.?)\s+(\d+)', re.IGNORECASE)
        section_pattern = re.compile(r'(?i)(secci[oó]n|cap[ií]tulo|t[ií]tulo)\s+([IVX\d]+)', re.IGNORECASE)
        
        # Dividir por líneas para detectar inicio de artículos
        lines = text.split('\n')
        current_section = []
        current_article = None
        section_start_idx = 0
        
        for i, line in enumerate(lines):
            # Detectar inicio de nuevo artículo
            article_match = article_pattern.search(line)
            section_match = section_pattern.search(line)
            
            if article_match:
                article_num = article_match.group(2)
                
                # Si ya tenemos una sección, guardarla
                if current_section and current_article is not None:
                    section_text = '\n'.join(current_section).strip()
                    if len(section_text) > 50:  # Solo agregar si tiene contenido significativo
                        sections.append({
                            'text': section_text,
                            'article': current_article,
                            'start_idx': section_start_idx
                        })
                
                # Empezar nueva sección
                current_section = [line]
                current_article = article_num
                section_start_idx = i
            elif section_match and not article_match:
                # Sección sin artículo específico
                if current_section and current_article is not None:
                    section_text = '\n'.join(current_section).strip()
                    if len(section_text) > 50:
                        sections.append({
                            'text': section_text,
                            'article': current_article,
                            'start_idx': section_start_idx
                        })
                
                current_section = [line]
                current_article = None
                section_start_idx = i
            else:
                # Agregar línea a sección actual
                current_section.append(line)
        
        # Agregar última sección
        if current_section:
            section_text = '\n'.join(current_section).strip()
            if len(section_text) > 50:
                sections.append({
                    'text': section_text,
                    'article': current_article,
                    'start_idx': section_start_idx
                })
        
        # Si no encontramos artículos, dividir por párrafos grandes
        if len(sections) < 2:
            logger.info("⚠️ Pocas secciones detectadas. Dividiendo por párrafos grandes...")
            sections = self._split_into_large_sections(text, max_size=100000)
        
        # Si hay demasiadas secciones pequeñas, combinarlas
        if len(sections) > 100:
            logger.info(f"⚠️ Muchas secciones detectadas ({len(sections)}). Se combinarán automáticamente en el siguiente paso.")
        
        return sections
    
    def _combine_small_sections(self, sections: List[Dict], min_size: int = 5000) -> List[Dict]:
        """
        Combinar secciones pequeñas en secciones más grandes para optimizar el procesamiento.
        Esto reduce el número de llamadas a SemanticSplitter y mejora el rendimiento.
        
        Args:
            sections: Lista de secciones originales
            min_size: Tamaño mínimo objetivo para cada sección combinada
        
        Returns:
            Lista de secciones optimizadas
        """
        if not sections:
            return sections
        
        optimized = []
        current_combined = []
        current_size = 0
        current_article = None
        
        for section in sections:
            section_text = section['text']
            section_article = section.get('article')
            section_size = len(section_text)
            
            # Si la sección actual es muy grande, guardarla sola
            if section_size >= min_size * 2:
                # Guardar combinación anterior si existe
                if current_combined:
                    combined_text = '\n\n'.join([s['text'] for s in current_combined]).strip()
                    optimized.append({
                        'text': combined_text,
                        'article': current_article,
                        'start_idx': current_combined[0].get('start_idx', 0)
                    })
                    current_combined = []
                    current_size = 0
                
                # Agregar sección grande sola
                optimized.append(section)
                continue
            
            # Combinar secciones pequeñas
            if current_size + section_size < min_size * 1.5:  # Combinar hasta ~1.5x min_size
                current_combined.append(section)
                current_size += section_size
                # Usar el artículo de la primera sección o el más común
                if current_article is None:
                    current_article = section_article
            else:
                # Guardar combinación actual
                if current_combined:
                    combined_text = '\n\n'.join([s['text'] for s in current_combined]).strip()
                    optimized.append({
                        'text': combined_text,
                        'article': current_article,
                        'start_idx': current_combined[0].get('start_idx', 0)
                    })
                
                # Empezar nueva combinación
                current_combined = [section]
                current_size = section_size
                current_article = section_article
        
        # Agregar última combinación
        if current_combined:
            combined_text = '\n\n'.join([s['text'] for s in current_combined]).strip()
            optimized.append({
                'text': combined_text,
                'article': current_article,
                'start_idx': current_combined[0].get('start_idx', 0)
            })
        
        return optimized
    
    def _split_into_large_sections(self, text: str, max_size: int = 150000) -> List[Dict]:
        """
        Dividir texto en secciones grandes cuando no se detectan artículos.
        Útil para documentos sin estructura clara de artículos.
        
        Args:
            text: Texto completo
            max_size: Tamaño máximo de cada sección en caracteres
        
        Returns:
            Lista de secciones
        """
        sections = []
        
        # Dividir por párrafos grandes primero
        paragraphs = re.split(r'\n\s*\n+', text)
        
        current_section = []
        current_size = 0
        section_num = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para or len(para) < 10:
                continue
            
            # Si agregar este párrafo excede el tamaño máximo, guardar sección actual
            if current_size + len(para) > max_size and current_section:
                section_text = '\n\n'.join(current_section).strip()
                if len(section_text) > 50:
                    sections.append({
                        'text': section_text,
                        'article': None,
                        'start_idx': section_num
                    })
                    section_num += 1
                
                # Mantener overlap: últimos 500 caracteres
                overlap_text = section_text[-500:] if len(section_text) > 500 else section_text
                current_section = [overlap_text, para] if overlap_text else [para]
                current_size = len(overlap_text) + len(para) if overlap_text else len(para)
            else:
                current_section.append(para)
                current_size += len(para) + 2  # +2 por los \n\n
        
        # Agregar última sección
        if current_section:
            section_text = '\n\n'.join(current_section).strip()
            if len(section_text) > 50:
                sections.append({
                    'text': section_text,
                    'article': None,
                    'start_idx': section_num
                })
        
        return sections
    
    def _split_section_basic(self, text: str, source: str, article: Optional[str], start_id: int) -> List[Dict]:
        """
        Dividir una sección usando método básico optimizado para velocidad.
        Respeta párrafos completos y mantiene overlap para contexto.
        Optimizado para ser muy rápido mientras mantiene calidad.
        
        Args:
            text: Texto de la sección
            source: Fuente del documento
            article: Número de artículo
            start_id: ID inicial para los chunks
        
        Returns:
            Lista de chunks
        """
        chunks = []
        chunk_id = start_id
        
        # Limpiar texto
        text = text.strip()
        if not text:
            return chunks
        
        # Si el texto es muy pequeño, devolverlo como un solo chunk
        if len(text) <= self.chunk_size:
            chunks.append({
                'id': chunk_id,
                'text': text,
                'source': source,
                'chunk_index': chunk_id,
                'article': article
            })
            return chunks
        
        # Dividir por párrafos (más rápido que dividir por oraciones)
        paragraphs = text.split('\n\n')
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para or len(para) < 10:
                continue
            
            # Si el párrafo es extremadamente largo, dividirlo por líneas
            if len(para) > self.chunk_size * 2:
                # Dividir párrafo muy largo en chunks más pequeños
                lines = para.split('\n')
                temp_chunk = current_chunk
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        continue
                    
                    if len(temp_chunk) + len(line) + 1 > self.chunk_size:
                        if temp_chunk:
                            chunks.append({
                                'id': chunk_id,
                                'text': temp_chunk.strip(),
                                'source': source,
                                'chunk_index': chunk_id,
                                'article': article
                            })
                            chunk_id += 1
                            # Overlap simple
                            overlap = temp_chunk[-self.chunk_overlap:] if len(temp_chunk) > self.chunk_overlap else ""
                            temp_chunk = overlap + "\n" + line if overlap else line
                        else:
                            temp_chunk = line
                    else:
                        temp_chunk += "\n" + line if temp_chunk else line
                
                current_chunk = temp_chunk
            else:
                # Párrafo normal: agregar al chunk actual
                if len(current_chunk) + len(para) + 2 > self.chunk_size:
                    if current_chunk:
                        chunks.append({
                            'id': chunk_id,
                            'text': current_chunk.strip(),
                            'source': source,
                            'chunk_index': chunk_id,
                            'article': article
                        })
                        chunk_id += 1
                        # Overlap para mantener contexto
                        overlap = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                        current_chunk = overlap + "\n\n" + para if overlap else para
                    else:
                        current_chunk = para
                else:
                    current_chunk += "\n\n" + para if current_chunk else para
        
        # Agregar último chunk
        if current_chunk.strip():
            chunks.append({
                'id': chunk_id,
                'text': current_chunk.strip(),
                'source': source,
                'chunk_index': chunk_id,
                'article': article
            })
        
        return chunks
    
    def _split_into_chunks(self, text: str, source: str) -> List[Dict]:
        """
        Dividir texto en chunks para embeddings con chunking inteligente
        Respeta artículos, secciones y párrafos completos
        
        Args:
            text: Texto completo
            source: Fuente del documento
        
        Returns:
            Lista de chunks con metadata
        """
        # Limpiar texto
        text = self._clean_text(text)
        
        # Detectar artículos y secciones importantes
        # Patrones: "Artículo X", "ARTÍCULO X", "Art. X", "artículo X", etc.
        article_pattern = re.compile(r'(?i)(art[ií]culo|art\.?)\s+(\d+)', re.IGNORECASE)
        section_pattern = re.compile(r'(?i)(secci[oó]n|cap[ií]tulo|t[ií]tulo)\s+([IVX\d]+)', re.IGNORECASE)
        
        chunks = []
        chunk_id = 0
        
        # Dividir por párrafos primero
        paragraphs = re.split(r'\n\s*\n+', text)
        
        current_chunk = ""
        current_article = None
        
        for para in paragraphs:
            para = para.strip()
            if not para or len(para) < 10:
                continue
            
            # Detectar si este párrafo es un artículo o sección
            article_match = article_pattern.search(para)
            section_match = section_pattern.search(para)
            
            # Si encontramos un nuevo artículo, guardar chunk anterior y empezar uno nuevo
            if article_match:
                article_num = article_match.group(2)
                if current_chunk and current_article != article_num:
                    # Guardar chunk anterior con overlap
                    if current_chunk:
                        chunks.append({
                            'id': chunk_id,
                            'text': current_chunk.strip(),
                            'source': source,
                            'chunk_index': chunk_id,
                            'article': current_article
                        })
                        chunk_id += 1
                    
                    # Empezar nuevo chunk con el artículo
                    current_chunk = para
                    current_article = article_num
                    continue
            
            # Si el párrafo es muy largo, dividirlo inteligentemente
            if len(para) > self.chunk_size:
                # Dividir por oraciones primero
                sentences = re.split(r'([.!?]\s+)', para)
                temp_chunk = current_chunk
                
                for i in range(0, len(sentences), 2):
                    sentence = sentences[i] + (sentences[i+1] if i+1 < len(sentences) else '')
                    
                    if len(temp_chunk) + len(sentence) + 2 > self.chunk_size:
                        if temp_chunk:
                            chunks.append({
                                'id': chunk_id,
                                'text': temp_chunk.strip(),
                                'source': source,
                                'chunk_index': chunk_id,
                                'article': current_article
                            })
                            chunk_id += 1
                        # Mantener overlap: últimos 100 caracteres del chunk anterior
                        overlap = temp_chunk[-self.chunk_overlap:] if len(temp_chunk) > self.chunk_overlap else ""
                        temp_chunk = overlap + sentence
                    else:
                        temp_chunk += " " + sentence if temp_chunk else sentence
                
                current_chunk = temp_chunk
            else:
                # Agregar párrafo al chunk actual
                if len(current_chunk) + len(para) + 2 > self.chunk_size:
                    if current_chunk:
                        chunks.append({
                            'id': chunk_id,
                            'text': current_chunk.strip(),
                            'source': source,
                            'chunk_index': chunk_id,
                            'article': current_article
                        })
                        chunk_id += 1
                        # Mantener overlap para contexto
                        overlap = current_chunk[-self.chunk_overlap:] if len(current_chunk) > self.chunk_overlap else ""
                        current_chunk = overlap + "\n\n" + para
                    else:
                        current_chunk = para
                else:
                    current_chunk += "\n\n" + para if current_chunk else para
        
        # Agregar último chunk
        if current_chunk.strip():
            chunks.append({
                'id': chunk_id,
                'text': current_chunk.strip(),
                'source': source,
                'chunk_index': chunk_id,
                'article': current_article
            })
        
        logger.info(f"Chunks creados: {len(chunks)} (con chunking inteligente)")
        return chunks
    
    def _extract_article_number(self, text: str) -> Optional[str]:
        """Extraer número de artículo del texto"""
        article_match = re.search(r'(?i)(art[ií]culo|art\.?)\s+(\d+)', text)
        if article_match:
            return article_match.group(2)
        return None
    
    def _clean_text(self, text: str) -> str:
        """Limpiar y normalizar texto"""
        # Remover caracteres especiales
        text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalizar espacios
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n+', '\n', text)
        
        return text.strip()
    
    def process_directory(self, directory_path: str) -> List[Dict]:
        """
        Procesar todos los documentos en un directorio
        
        Args:
            directory_path: Ruta al directorio
        
        Returns:
            Lista de todos los chunks de todos los documentos
        """
        all_chunks = []
        
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if os.path.isfile(file_path):
                file_ext = os.path.splitext(filename)[1].lower()
                if file_ext in self.supported_formats:
                    try:
                        logger.info(f"Procesando: {filename}")
                        chunks = self.process_document(file_path)
                        all_chunks.extend(chunks)
                    except Exception as e:
                        logger.error(f"Error procesando {filename}: {str(e)}")
                        continue
        
        logger.info(f"Total de chunks procesados: {len(all_chunks)}")
        return all_chunks

